import os
import re
import json
import asyncio
from datetime import datetime, timedelta
from collections import defaultdict
from docx import Document
from app.database import db

SCHEDULE_FILE = "Расписание.docx"
SHIFTS_FILE = "group_shifts.json"

days_ru = ["Понедельник", "Вторник", "Среда", "Четверг", "Пятница", "Суббота"]

# ============== ФУНКЦИИ ПАРСИНГА ==============


def normalize_teacher_name(name: str):
    """Нормализует имя преподавателя в формат 'Фамилия И.О.'"""
    if not name:
        return None
    name = re.sub(r"\s+", " ", name.strip())
    m = re.search(r"([А-ЯЁ][а-яё]+(?:-[А-ЯЁ][а-яё]+)?\s+[А-ЯЁ]\.[А-ЯЁ]\.?)", name)
    if m:
        name = m.group(1)
    name = re.sub(r"\s*\d{2,4}[А-Яа-я]?$", "", name)
    name = re.sub(r"([А-ЯЁ])\.([А-ЯЁ])$", r"\1.\2.", name)
    name = re.sub(r"\.\.", ".", name)
    name = re.sub(r"[^А-Яа-яЁё.\s-]", "", name).strip()
    if not re.match(r"^[А-ЯЁ][а-яё-]+\s+[А-ЯЁ]\.[А-ЯЁ]\.?$", name):
        return None
    fam, ini = name.split(maxsplit=1)
    fam = "-".join(s[:1].upper() + s[1:].lower() for s in fam.split("-"))
    return f"{fam} {ini}"


def parse_lesson_info_fixed(cell_text: str):
    """
    Парсит предмет, преподавателя и кабинет из ячейки DOCX.
    Учитывает:
    - МДК 07.01 ...
    - кабинет рядом с преподавателем
    """

    if not cell_text:
        return None

    # сохраняем переносы строк
    lines = [l.strip() for l in cell_text.split("\n") if l.strip()]
    text = " ".join(lines)

    if not text or text == "##":
        return None

    teacher = None
    classroom = None

    # --- ищем преподавателя ---
    teacher_match = re.search(
        r"([А-ЯЁ][а-яё]+(?:-[А-ЯЁ][а-яё]+)?\s+[А-ЯЁ]\.[А-ЯЁ]\.?)", text
    )

    if teacher_match:
        teacher = teacher_match.group(1)

        # ищем кабинет рядом с преподавателем
        after_teacher = text[teacher_match.end() :]

        room_match = re.search(r"\b\d{2,3}\b", after_teacher)
        if room_match:
            classroom = room_match.group(0)

    # --- предмет ---
    subject = text

    if teacher:
        subject = subject.replace(teacher, "")

    if classroom:
        subject = subject.replace(classroom, "")

    subject = re.sub(r"\s+", " ", subject).strip()

    # если кабинет всё ещё внутри subject — убираем
    if subject and classroom:
        subject = re.sub(rf"\b{classroom}\b", "", subject).strip()

    if not teacher:
        print("⚠ Не найден преподаватель:", cell_text)

    return {
        "subject": subject if subject else None,
        "teacher": teacher,
        "classroom": classroom,
    }


def add_classrooms_to_schedule(schedule: dict, group_name: str, shifts: dict):
    """
    Добавляет кабинеты из group_shifts.json,
    НО только если они отсутствуют в расписании.
    """

    group_shift = shifts.get(group_name, {})
    classroom = group_shift.get("room")

    if not classroom:
        return schedule

    for day in schedule["zero_lesson"]:
        lesson = schedule["zero_lesson"][day]
        if lesson and not lesson.get("classroom"):
            lesson["classroom"] = classroom

    for day in schedule["days"]:
        for lesson_num, lesson in schedule["days"][day].items():
            if lesson and not lesson.get("classroom"):
                lesson["classroom"] = classroom

    return schedule


def parse_schedule_table_fixed(table, group_name: str, schedules: dict):
    """Парсит таблицу с расписанием с поддержкой подпар (половинок)"""
    print(f"\n--- [ТАБЛИЦА] Начало парсинга для группы: {group_name} ---")
    rows = table.rows
    print(f"[ТАБЛИЦА] Всего строк в таблице: {len(rows)}")

    if not rows:
        print("[ТАБЛИЦА] Строк нет, выход.")
        return

    # 1. Логирование заголовков
    header = [cell.text.strip() for cell in rows[0].cells]
    print(f"[ЗАГОЛОВОК] Ячейки заголовка: {header}")

    if len(header) < 2:
        print("[ЗАГОЛОВОК] Слишком мало колонок, пропускаем.")
        return

    # 2. Логирование определения дней недели
    day_columns = {}
    for idx, cell in enumerate(header[1:], 1):
        for day in days_ru:
            if day in cell:
                day_columns[idx] = day
                print(f"[ДНИ] Колонка {idx} определена как: {day} (текст: '{cell}')")
                break
        if idx not in day_columns:
            print(f"[ДНИ] Колонка {idx} НЕ распознана как день (текст: '{cell}')")

    print(f"[ДНИ] Итоговая карта дней: {day_columns}")

    # === ПЕРВЫЙ ПРОХОД: Собираем все данные временно ===
    temp_schedule = defaultdict(lambda: defaultdict(list))

    # ИСПРАВЛЕНИЕ: Теперь храним СПИСОК записей для каждого signature
    # (day, subject, teacher, classroom) -> [(row_idx, lesson_num, lesson_info, is_merged), ...]
    seen_cells = defaultdict(list)

    # Отслеживаем ТОЛЬКО реально объединённые ячейки
    merged_cells = set()

    for r_idx, r in enumerate(rows[1:], 1):
        cells = [c.text.strip() for c in r.cells]

        if not cells or not cells[0]:
            continue

        lesson_num = cells[0].strip()
        if not re.match(r"^\d+$", lesson_num):
            continue

        print(f"[СТРОКА {r_idx}] Номер пары: {lesson_num}")

        for idx, text in enumerate(cells[1:], 1):
            if idx not in day_columns:
                continue

            if not text.strip():
                continue

            day = day_columns[idx]
            print(f"[ЯЧЕЙКА {r_idx},{idx}] День: {day}, Текст: '{text}'")

            lesson_info = parse_lesson_info_fixed(text)

            if lesson_info:
                print(f"  -> [УРОК] Распознано: {lesson_info}")

                cell_signature = (
                    day,
                    lesson_info.get("subject"),
                    lesson_info.get("teacher"),
                    lesson_info.get("classroom"),
                )

                # Проверяем есть ли уже записи с этим signature
                existing_entries = seen_cells[cell_signature]

                if existing_entries:
                    # Берем последнюю запись для этого signature
                    prev_row_idx, prev_lesson_num, _, prev_is_merged = existing_entries[
                        -1
                    ]

                    # ИСПРАВЛЕНИЕ: Объединённая ячейка только если номер пары НЕ изменился
                    if int(lesson_num) == int(prev_lesson_num):
                        # Одинаковые номера = объединённая ячейка
                        print(
                            f"  [ОБЪЕДИНЁННАЯ] Ячейка охватывает строки {prev_row_idx}-{r_idx}, номер пары {lesson_num}"
                        )
                        merged_cells.add(cell_signature)
                        # Обновляем последнюю запись
                        existing_entries[-1] = (r_idx, lesson_num, lesson_info, True)
                    # ИСПРАВЛЕНИЕ: Если номер пары изменился - это разные занятия!
                    elif int(lesson_num) > int(prev_lesson_num):
                        print(
                            f"  [РАЗНЫЕ ПАРЫ] Номер изменился {prev_lesson_num} -> {lesson_num}, добавляем новую запись"
                        )
                        # ДОБАВЛЯЕМ новую запись, не перезаписываем!
                        existing_entries.append((r_idx, lesson_num, lesson_info, False))
                    else:
                        print(
                            f"  [ПРОПУСК] Дубликат (номер {lesson_num} < {prev_lesson_num})"
                        )
                else:
                    # Первая запись для этого signature
                    seen_cells[cell_signature].append(
                        (r_idx, lesson_num, lesson_info, False)
                    )

    # Переносим в temp_schedule
    for cell_signature, entries in seen_cells.items():
        day, subject, teacher, classroom = cell_signature

        # is_merged=True ТОЛЬКО если ячейка была в merged_cells
        is_merged = cell_signature in merged_cells

        for row_idx, lesson_num, lesson_info, _ in entries:
            temp_schedule[day][lesson_num].append((row_idx, lesson_info, is_merged))
            print(f"[TEMP] {day} пара {lesson_num}: {subject} (merged={is_merged})")

    # === ВТОРОЙ ПРОХОД: Определяем подпары и сохраняем ===
    for day, lessons in temp_schedule.items():
        for lesson_num, lesson_list in lessons.items():
            if lesson_num == "0":
                if lesson_list:
                    schedules[group_name]["zero_lesson"][day] = lesson_list[0][1]
                continue

            lesson_list.sort(key=lambda x: x[0])

            print(f"\n[АНАЛИЗ] {day}, пара {lesson_num}: {len(lesson_list)} записей")

            if len(lesson_list) == 1:
                row_idx, lesson_info, is_merged = lesson_list[0]

                # Если это объединённая ячейка - она ЦЕЛАЯ
                if is_merged:
                    lesson_key = lesson_num
                    print(f"  [ЦЕЛАЯ из merged] {day} {lesson_num} -> {lesson_key}")
                    schedules[group_name]["days"][day][lesson_key] = lesson_info
                else:
                    # Проверяем, есть ли другие дни с двумя записями
                    has_second_half = False
                    for other_day, other_lessons in temp_schedule.items():
                        if other_day == day:
                            continue
                        if (
                            lesson_num in other_lessons
                            and len(other_lessons[lesson_num]) > 1
                        ):
                            has_second_half = True
                            break

                    if has_second_half:
                        lesson_key = f"{lesson_num}.1"
                        print(f"  [ПОЛОВИНКА] {day} {lesson_num} -> {lesson_key}")
                    else:
                        lesson_key = lesson_num
                        print(f"  [ЦЕЛАЯ] {day} {lesson_num} -> {lesson_key}")

                    schedules[group_name]["days"][day][lesson_key] = lesson_info

            else:
                # Несколько записей в одном дне - подпары
                for i, (row_idx, lesson_info, is_merged) in enumerate(lesson_list, 1):
                    lesson_key = f"{lesson_num}.{i}"

                    print(f"  [ПОДПАРА] {day} {lesson_key} (строка {row_idx})")
                    schedules[group_name]["days"][day][lesson_key] = lesson_info

    print(f"--- [ТАБЛИЦА] Конец парсинга для группы: {group_name} ---\n")


def parse_schedule_from_docx(file_path: str):
    """Парсит расписание из DOCX и возвращает dict с логом"""
    print(f"=== [DOCX] Открытие файла: {file_path} ===")
    doc = Document(file_path)
    schedules = {}
    current_group = None

    # 1. Поиск групп в параграфах
    print(f"[DOCX] Всего параграфов: {len(doc.paragraphs)}")
    for p_idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue

        group_match = re.search(r"Расписание уроков\s+для\s+(.+?)\s+группы", text)
        if group_match:
            current_group = group_match.group(1).strip()
            print(f"[ГРУППА] Найдена группа '{current_group}' в параграфе {p_idx}")
            schedules[current_group] = {
                "days": {day: {} for day in days_ru},
                "zero_lesson": {day: {} for day in days_ru},
            }
            continue

    print(f"[DOCX] Итого найдено групп: {list(schedules.keys())}")

    # 2. Поиск таблиц
    # Фильтруем таблицы, где есть упоминание дней недели
    tables = []
    for t_idx, t in enumerate(doc.tables):
        # Собираем весь текст таблицы для проверки
        table_text = "\n".join(cell.text for row in t.rows for cell in row.cells)
        has_day = any(day in table_text for day in days_ru)
        if has_day:
            tables.append(t)
            print(
                f"[ТАБЛИЦА #{t_idx}] Найдена таблица с расписанием (содержит дни недели)"
            )
        else:
            # print(f"[ТАБЛИЦА #{t_idx}] Пропущена (нет дней недели)")
            pass

    print(f"[DOCX] Найдено подходящих таблиц: {len(tables)}")

    # 3. Сопоставление таблиц и групп
    group_index = 0
    group_names = list(schedules.keys())

    print(f"[СОПОСТАВЛЕНИЕ] Групп: {len(group_names)}, Таблиц: {len(tables)}")

    for table_idx, table in enumerate(tables):
        if group_index >= len(group_names):
            print(
                f"[ОШИБКА] Таблиц больше, чем групп. Остановка на таблице {table_idx}"
            )
            break

        group_name = group_names[group_index]
        print(f"[СОПОСТАВЛЕНИЕ] Таблица #{table_idx} -> Группа '{group_name}'")

        parse_schedule_table_fixed(table, group_name, schedules)
        group_index += 1

    print("=== [DOCX] Парсинг завершен ===")
    return schedules


# ============== РАБОТА СО СМЕНАМИ И БД ==============


def load_group_shifts():
    if not os.path.exists(SHIFTS_FILE):
        return {}
    try:
        with open(SHIFTS_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
            normalized = {}
            for k, v in data.items():
                if isinstance(v, dict):
                    normalized[k] = {
                        "shift": v.get("shift", 1),
                        "room": v.get("room", ""),
                    }
                else:
                    normalized[k] = {"shift": v, "room": ""}
            return normalized
    except Exception as e:
        print(f"Ошибка загрузки смен: {e}")
        return {}


# ============== АСИНХРОННЫЙ ПЛАНИРОВЩИК ==============


async def load_schedule_to_db():
    """Парсит и сохраняет расписание в MongoDB"""
    if not os.path.exists(SCHEDULE_FILE):
        print(f"❌ Файл {SCHEDULE_FILE} не найден")
        return
    data = parse_schedule_from_docx(SCHEDULE_FILE)
    if not data:
        print("❌ Не удалось распарсить расписание.")
        return

    # Загрузка смен
    shifts = load_group_shifts()

    # Очистим старые записи
    await db.schedules.delete_many({})
    for group, schedule in data.items():
        # Добавляем кабинеты из group_shifts.json
        schedule_with_classrooms = add_classrooms_to_schedule(schedule, group, shifts)

        await db.schedules.insert_one(
            {
                "group_name": group,
                "schedule": schedule_with_classrooms,
                "shift_info": shifts.get(group, {"shift": 1}),
                "updated_at": datetime.now(),
            }
        )
    print(f"✅ Залито расписание для {len(data)} групп в MongoDB")
